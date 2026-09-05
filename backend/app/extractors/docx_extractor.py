import docx
from typing import Dict, Any, List

def extract_docx(file_path: str) -> Dict[str, Any]:
    pages = []
    full_text_parts = []
    current_section_text = []
    section_count = 1
    
    try:
        doc = docx.Document(file_path)
        
        # Process paragraphs and headers
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
                
            if p.style and 'Heading' in p.style.name:
                if current_section_text:
                    sec_content = "\n".join(current_section_text)
                    pages.append({"page_number": section_count, "text": sec_content})
                    full_text_parts.append(f"--- Section {section_count} ---\n{sec_content}")
                    section_count += 1
                    current_section_text = []
                current_section_text.append(f"## {text}")
            else:
                current_section_text.append(text)
                
        # Process tables inside Word doc
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_rows.append(" | ".join(row_data))
            if table_rows:
                current_section_text.append("\n[Word Table]:\n" + "\n".join(table_rows))
                
        if current_section_text:
            sec_content = "\n".join(current_section_text)
            pages.append({"page_number": section_count, "text": sec_content})
            full_text_parts.append(f"--- Section {section_count} ---\n{sec_content}")
            
    except Exception as e:
        raise ValueError(f"Failed to extract Word document content: {str(e)}")

    full_text = "\n\n".join(full_text_parts) if full_text_parts else "Empty Word Document"
    
    return {
        "file_type": "docx",
        "page_count": max(1, len(pages)),
        "full_text": full_text,
        "pages": pages if pages else [{"page_number": 1, "text": full_text}],
        "is_tabular": False,
        "tabular_stats": {}
    }
